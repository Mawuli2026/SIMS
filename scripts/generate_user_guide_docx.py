from __future__ import annotations

import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "USER_GUIDE.md"
OUTPUT = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else ROOT / "SIMS_User_Guide.docx"
LOGO = ROOT / "public" / "sims-logo.png"

NAVY = "17365D"
BLUE = "2563EB"
LIGHT_BLUE = "EAF2FF"
LIGHT_GRAY = "F3F4F6"
WHITE = "FFFFFF"
TEXT = RGBColor(31, 41, 55)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def add_inline_markdown(paragraph, text: str) -> None:
    token_pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^]]+\]\([^)]+\))")
    position = 0
    for match in token_pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(30, 64, 175)
        else:
            link_match = re.match(r"\[([^]]+)\]\(([^)]+)\)", token)
            paragraph.add_run(link_match.group(1) if link_match else token)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(32)
    title.font.bold = True
    title.font.color.rgb = RGBColor(23, 54, 93)

    for style_name, size, color in (
        ("Heading 1", 18, NAVY),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 11.5, NAVY),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    if "Guide Subtitle" not in styles:
        subtitle = styles.add_style("Guide Subtitle", WD_STYLE_TYPE.PARAGRAPH)
        subtitle.font.name = "Aptos"
        subtitle.font.size = Pt(15)
        subtitle.font.color.rgb = RGBColor(75, 85, 99)
        subtitle.paragraph_format.space_after = Pt(10)


def configure_sections(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)

        header = section.header.paragraphs[0]
        header.text = "SIMS  |  User Guide"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.runs[0].font.name = "Aptos"
        header.runs[0].font.size = Pt(8.5)
        header.runs[0].font.color.rgb = RGBColor(107, 114, 128)

        footer = section.footer.paragraphs[0]
        add_page_number(footer)
        for run in footer.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(107, 114, 128)


def add_title_page(document: Document) -> None:
    if LOGO.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(42)
        paragraph.add_run().add_picture(str(LOGO), width=Inches(1.55))

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(20)
    title.add_run("SIMS User Guide")

    subtitle = document.add_paragraph(style="Guide Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Sales and Inventory Management System")

    audience = document.add_paragraph()
    audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
    audience.paragraph_format.space_before = Pt(20)
    audience.add_run("Operating instructions for\nSystemAdmin, Manager, and Cashier").bold = True

    version = document.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.paragraph_format.space_before = Pt(85)
    version.add_run(f"Version 1.1  |  {date.today().strftime('%B %Y')}")

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(8)
    note.add_run("School Project — Lite Edition").italic = True
    document.add_page_break()


def add_table(document: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for row_index, source_row in enumerate(rows):
        for column_index, value in enumerate(source_row):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_markdown(paragraph, value)
            if row_index == 0:
                set_cell_shading(cell, NAVY)
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines: list[str] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        table_lines.append(lines[index].strip())
        index += 1
    parsed = [[cell.strip() for cell in line.strip("|").split("|")] for line in table_lines]
    if len(parsed) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in parsed[1]):
        parsed.pop(1)
    return parsed, index


def add_markdown_content(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 1  # The first Markdown title is represented by the Word title page.
    first_section = True

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()
        if not line:
            index += 1
            continue

        if line.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1)) - 1
            if level == 1 and not first_section:
                document.add_page_break()
            paragraph = document.add_heading(level=level)
            add_inline_markdown(paragraph, heading_match.group(2))
            if level == 1:
                first_section = False
            index += 1
            continue

        number_match = re.match(r"^\d+\.\s+(.+)$", line)
        if number_match:
            paragraph = document.add_paragraph(style="List Number")
            add_inline_markdown(paragraph, number_match.group(1))
            index += 1
            continue

        bullet_match = re.match(r"^-\s+(.+)$", line)
        if bullet_match:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, bullet_match.group(1))
            index += 1
            continue

        paragraph = document.add_paragraph()
        add_inline_markdown(paragraph, line)
        index += 1


def build_document() -> None:
    document = Document()
    document.core_properties.title = "SIMS User Guide"
    document.core_properties.subject = "Role-based operating instructions for the Sales and Inventory Management System"
    document.core_properties.author = "SIMS Development Team"
    document.core_properties.keywords = "SIMS, sales, inventory, SystemAdmin, Manager, Cashier"

    configure_styles(document)
    configure_sections(document)
    add_title_page(document)
    add_markdown_content(document, SOURCE.read_text(encoding="utf-8"))

    # Reapply header/footer settings after content-created page boundaries.
    configure_sections(document)
    document.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    build_document()
